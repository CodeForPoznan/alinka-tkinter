from docx.shared import Pt
from docx.enum.text import WD_BREAK

import random


def normal_center(document, text, size=10, bold=False):
    paragraph = document.add_paragraph(style='Normal_center_my')

    run = paragraph.add_run(text)
    run.font.bold = bold
    run.font.size = Pt(size)


def normal_right(document, text, size=10, bold=False):
    paragraph = document.add_paragraph(style='Normal_right_my')

    run = paragraph.add_run(text)
    run.font.bold = bold
    run.font.size = Pt(size)


def normal_left(document, text, size=10, bold=False):
    paragraph = document.add_paragraph(style='Normal_left_my')

    run = paragraph.add_run(text)
    run.font.bold = bold
    run.font.size = Pt(size)


def normal_justify(document, text, size=10, bold=False):
    paragraph = document.add_paragraph(style='Normal_justify_my')

    run = paragraph.add_run(text)
    run.font.bold = bold
    run.font.size = Pt(size)


def normal_bold(document, text):
    document.add_paragraph(
        text=text,
        style='Normal_bold_my'
    )


def add_line(document, width):
    document.add_paragraph(text='', style='Normal_center_my').size = Pt(width)


def staff(document, value):
    document.add_paragraph(
        'Zespół Orzekający przy Poradni Psychologiczno-Pedagogicznej'
        ' w Grodzisku Wlkp.\nul. Zbąszyńska 11\n',
        style='Header_my'
    )
    document.add_paragraph(
        'w składzie:',
        style='Small_my'
    )
    for i in value['staff']['team']:
        document.add_paragraph(
            '{} - {}'.format(i[0], i[1]),
            style='Small_my'
        )


def staff_decree(document, value):
    document.add_paragraph(
        'Skład zespołu orzekającego:\n',
        style='Normal_left_my'
    )
    for i in value['staff']['team']:
        document.add_paragraph(
            '{} - {}'.format(i[0], i[1]),
            style='Normal_left_my'
        )

def find_staff(values, specialization):
    list_of_staff = []
    for specialist in values['staff']['team']:
        if specialization in specialist[1]:
            list_of_staff.append(specialist)
    if specialization == 'dowolny':
        for specialist in values['staff']['team']:
            if 'zespołu' not in specialist[1]:
                list_of_staff.append(specialist)
    if not list_of_staff:
        list_of_staff = values['staff']['team']

    choosen = random.choice(list_of_staff)
    name = choosen[0].split(' ')[1]
    if name[-1] == 'a':
        sex = False
    else:
        sex = True
    return [choosen, sex]

def referent_speech(data):
    if data[1]:
        speech = 'powiedział'
    else:
        speech = 'powiedziała'
    return '{} - {} {}: '.format(
        data[0][0],
        data[0][1],
        speech
    )


def referent_speech1(data):
    if data[1]:
        speech = 'zapoznał'
    else:
        speech = 'zapoznała'
    return '{} - {} {} '.format(
        data[0][0],
        data[0][1],
        speech
    )


def text_with_tag(document, text, tag):
    """add line with text"""
    line1 = document.add_paragraph()
    line1.style = document.styles['Text_my']
    line1.add_run(text)

    # add line with tag

    line2 = document.add_paragraph()
    line2.style = document.styles['Tag_my']
    line2.add_run(tag)


def applicant(document, value):
    line1 = document.add_paragraph()
    line1.style = document.styles['Text_my']
    line1.add_run('\n Na wniosek: ').bold = False
    line1.add_run('{}'.format(value))

    line2 = document.add_paragraph('      (imię i nazwisko wnioskodawcy)\n')
    line2.style = document.styles['Tag_my']


def line_with_bold(documnet, text1, text2):
    line = documnet.add_paragraph(style='Normal_left_my')
    line.add_run(text1)
    line.add_run(text2).bold = True


def get_underlined(reason):
    list_of_underlined = []
    list_of_underlined.append(False)
    list_of_underlined.append(reason == 'niesłyszenie')
    list_of_underlined.append(reason == 'słabosłyszenie')
    list_of_underlined.append(reason == 'niewidzenie')
    list_of_underlined.append(reason == 'słabowidzenie')
    list_of_underlined.append(
        reason == 'niepełnosprawność ruchową, w tym z afazją'
    )
    list_of_underlined.append(
        reason == 'niepełnosprawność intelektualną w stopniu lekkim'
    )
    list_of_underlined.append(
        reason == 'niepełnosprawność intelektualną w stopniu umiarkowanym'
    )
    list_of_underlined.append(
        reason == 'niepełnosprawność intelektualną w stopniu znacznym'
    )
    list_of_underlined.append(reason == 'autyzm')
    list_of_underlined.append(reason == 'sprzężona')
    list_of_underlined.append(
        reason == 'zagrożenie niedostosowaniem społecznym'
    )
    list_of_underlined.append(reason == 'niedostosowanie społeczne')

    return list_of_underlined


def reason(document, value):
    reasons_list = (
        '1) niepełnosprawność dziecka lub ucznia:',
        '   a) niesłyszące',
        '   b) słabosłyszące',
        '   c) niewidzące',
        '   d) słabowidzące',
        '   e) niepełnosprawne ruchowo, w tym z afazją',
        '   f) niepełnosprawne intelektualnie w stopniu lekkim',
        '   g) niepełnosprawne intelektualnie w stopniu umiarkowanym',
        '   h) niepełnosprawne intelektualnie w stopniu znacznym',
        '   i) z autyzmem, w tym z zespołem Aspergera',
        '   j) z niepełnosprawnością sprzężoną&: ',
        '2) niedostosowanie społeczne',
        '3) zagrożenie niedostosowaniem społecznym.'
    )

    underlined_list = get_underlined(value['reason'][0])
    document.add_paragraph('ze względu na&:', style='Normal_left_my')

    next_lines = document.add_paragraph(style='Normal_left_my')
    for i in range(0, len(reasons_list)):
        if (
            value['reason'][0] == 'sprzężona' and
            reasons_list[i] == '   j) z niepełnosprawnością sprzężoną&: '
        ):
            next_lines.add_run('   j) ')
            next_lines.add_run(
                'z niepełnosprawnością sprzężoną&: {} i {}\n'.format(
                    value['reason'][1],
                    value['reason'][2]
                )
            ).font.underline = True
        else:
            place_bracket = reasons_list[i].find(')') + 1
            un_underlined_part = reasons_list[i][0:place_bracket + 1]
            underlined_part = reasons_list[i][place_bracket + 1:]
            next_lines.add_run(un_underlined_part)
            next_lines.add_run(
                '{}\n'.format(underlined_part)
            ).font.underline = underlined_list[i]


def reason_individual_preschool(document, reason):
    if 'uniemożliwiający' in reason:
        unable = False
        barely_can = True
    else:
        unable = True
        barely_can = False
    paragraph = document.add_paragraph(style='Normal_left_my')
    paragraph.add_run('ze względu na stan zdrowia dziecka ')
    paragraph.add_run('uniemożliwiający').font.strike = unable
    paragraph.add_run('/ ')
    paragraph.add_run('znacznie utrudniający').font.strike = barely_can
    paragraph.add_run(
        '& uczęszczanie do przedszkola, oddziału przedszkolnego w szkole'
        ' podstawowej lub innej formy wychowania przedszkolnego.'
    ).font.strike = False


def reason_individual(document, reason):
    if 'uniemożliwiający' in reason:
        unable = False
        barely_can = True
    else:
        unable = True
        barely_can = False
    paragraph = document.add_paragraph(style='Normal_left_my')
    paragraph.add_run('ze względu na stan zdrowia ucznia ')
    paragraph.add_run('uniemożliwiający').font.strike = unable
    paragraph.add_run('/ ')
    paragraph.add_run('znacznie utrudniający').font.strike = barely_can
    paragraph.add_run('& uczęszczanie do szkoły.').font.strike = False


def page_break(document):
    line1 = document.add_paragraph()
    line1.add_run().add_break(WD_BREAK.PAGE)


def recommendations(document, recommendations):
    number = 1
    for recommend in recommendations:
        line = document.add_paragraph(style='Normal_justify_my')
        line.add_run('{}) '.format(number)).font.size = Pt(10)
        line.add_run(recommend).font.size = Pt(8)
        document.add_paragraph('', style='Normal_left_my')
        number += 1


def insert_reciver(document, values):
    line = document.add_paragraph(style='Normal_left_my')
    reciver = line.add_run('Otrzymuje:')
    reciver.font.underline = True
    reciver.font.bold = True
    document.add_paragraph('Wnioskodawca:', style='Normal_left_my')
    document.add_paragraph(
        values['applicant_n'],
        style='Normal_left_my'
    )
    sub = document.add_paragraph(style='Normal_left_my')
    sub.add_run('                    (imię i nazwisko)').font.size = Pt(8)
    document.add_paragraph(
        '{}, {} {}'.format(
            values['address'],
            values['zip_code'],
            values['city']
        ),
        style='Normal_left_my')
    sub = document.add_paragraph(style='Normal_left_my')
    sub.add_run(
        '       (adres zamieszkania albo adres do korespondencji\n'
        '          – jeżeli jest inny niż adres zamieszkania)'
    ).font.size = Pt(8)


def application_subject(values):
    if values['subject'] == 'kształcenie specjalne':
        return 'orzeczenia o potrzebie kształcenia specjalnego'
    elif values['subject'] == 'indywidualne roczne przygotowanie przedszkolne':
        return (
            'orzeczenia o potrzebie indywidualnego '
            'rocznego przygotowania przedszkolnego'
        )
    elif values['subject'] == 'indywidualne nauczanie':
        return 'orzeczenia o potrzebie indywidualnego nauczania'
    elif values['subject'] == 'wczesne wspomaganie rozwoju':
        return 'opinii o potrzebie wczesnego wspomagania rozwoju'
    elif values['subject'] == 'zajęcia rewalidacyjno-wychowawcze indywidualne':
        return (
            'orzeczenia o potrzebie zajęć rewalidacyjno - '
            'wychowawczych indywidualnych'
        )
    elif values['subject'] == 'zajęcia rewalidacyjno-wychowawcze zespołowe':
        return (
            'orzeczenia o potrzebie zajęć rewalidacyjno - '
            'wychowawczych zespołowych'
        )


def application_reason(reason):
    if reason[0] == 'sprzężona':
        return 'niepełnosprawność sprzężoną: {} i {}'.format(
            reason[1],
            reason[2]
        )
    else:
        return reason[0]


def individual_group(document, values):
    paragraph = document.add_paragraph(style='Normal_center_my')
    paragraph.add_run(
        'orzeka o potrzebie zajęć rewalidacyjno-wychowawczych '
    ).bold = True
    if values['subject'] == 'zajęcia rewalidacyjno-wychowawcze zespołowe':
        run = paragraph.add_run('zespołowych')
        run.bold = True
        run.underline = True
        paragraph.add_run('/indywidualnych.&').bold = True
    else:
        paragraph.add_run('zespołowych/').bold = True
        run = paragraph.add_run('indywidualnych')
        run.bold = True
        run.underline = True
        run1 = paragraph.add_run('.&')
        run1.bold = True
        run1.underline = False
